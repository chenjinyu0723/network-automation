# ruff: noqa: E501
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "工业交换机自动配置用户手册.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "183146"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
CALLOUT = "F4F6F9"
TABLE_WIDTH = 9360


def set_font(run, name: str = "Calibri", size: float = 11, *, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)


def add_page_field(paragraph) -> None:
    paragraph.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    paragraph.add_run(" 页")


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)
    set_keep_with_next(paragraph)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_font(rest)
    else:
        run = paragraph.add_run(text)
        set_font(run)


def add_list(doc: Document, items: list[str], *, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        run = paragraph.add_run(item)
        set_font(run)


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.16)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F8FA")
    p_pr.append(shd)
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_font(run, "Consolas", 9.5, color="244052")


def add_callout(doc: Document, title: str, text: str, *, caution: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FFF5E6" if caution else CALLOUT)
    p_pr.append(shd)
    heading = paragraph.add_run(title + "  ")
    set_font(heading, size=10.5, bold=True, color="7A5A00" if caution else DARK_BLUE)
    body = paragraph.add_run(text)
    set_font(body, size=10.5, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for cell, value in zip(header_cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        set_font(run, size=10, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])
    for source_row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, source_row):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            set_font(run, size=10)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in ((1, 16, BLUE, 18, 10), (2, 13, BLUE, 14, 7), (3, 12, DARK_BLUE, 10, 5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("AI Agent 工业交换机自动配置 | 用户手册")
    set_font(header_run, size=9, color="60788A")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(footer)
    for run in footer.runs:
        set_font(run, size=9, color="60788A")


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("AI Agent 工业交换机自动配置")
    set_font(r, "Microsoft YaHei", 25, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("用户手册")
    set_font(r, "Microsoft YaHei", 19, bold=True, color=BLUE)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["适用版本", "当前本地单用户桌面版"],
            ["更新时间", "2026-08-20"],
            ["运行范围", "Windows 本机，界面仅访问 127.0.0.1"],
            ["项目地址", "https://github.com/chenjinyu0723/network-automation"],
            ["许可证", "MIT License；厂商手册、用户数据和设备配置不因该许可证获得额外授权。"],
        ],
        [2200, 7160],
    )
    add_callout(
        doc,
        "使用边界",
        "应用生成的是可编辑的配置思路和命令草案。是否下发、下发到哪台设备、以及现场变更后的验收，始终由具备授权的操作人员决定。",
        caution=True,
    )


def build_manual() -> None:
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_heading(doc, "目录", 1)
    add_body(doc, "可在 Word 左侧“导航”窗格按标题跳转；以下目录与正文标题一一对应。")
    add_list(doc, [
        "1. 应用概览与边界",
        "2. 启动前准备",
        "3. 启动、关闭与本地数据",
        "4. 首次体验：完成一次三交换机组网",
        "5. 设置：LLM 与 Embedding",
        "6. 手册管理：导入、抽取与向量索引",
        "7. 拓扑编辑：设备、直线链路与接口",
        "8. 配置规划：思路、检索与逐设备命令",
        "9. 模板管理：保存、查看与迁移",
        "10. 下发与结果：逐台执行与设备回显",
        "11. 导入、导出、删除与备份",
        "12. 常见问题与排查",
        "13. GitHub、许可证与二次开发",
        "附录：操作前检查清单",
    ], numbered=True)

    add_heading(doc, "1. 应用概览与边界", 1)
    add_body(doc, "本应用是一个在 Windows 本机运行的工业交换机自动配置工作区。用户先导入厂商手册、绘制真实拓扑，再由大模型起草配置思路。用户可以直接修改思路；确认非空思路后，系统检索本次选择的手册并为每台交换机生成可编辑命令。最后由用户逐台确认，通过 SSH 下发。")
    add_table(doc, ["能力", "当前行为"], [
        ["手册", "支持 CHM、PDF、HTML/HTM、TXT 与 Markdown。每个任务只选择一份已完成抽取的手册作为本次命令检索上下文。"],
        ["拓扑", "完全手工绘制，不做 LLDP/CDP 自动发现。链路保存两端真实接口名。"],
        ["规划", "先生成可编辑配置思路，再用两轮手册检索和逐设备命令生成形成草案。"],
        ["模板", "保存拓扑、需求、思路和最终命令，便于查看与迁移；模板不会参与新任务的提示词或命令生成。"],
        ["执行", "只允许用户逐台确认。系统不做一键批量下发；PC 不提供 SSH 登录或 ping 验收功能。"],
    ], [2100, 7260])
    add_callout(doc, "重要", "应用不会因为能力不在预置场景中就拒绝生成。手册证据、模型推断和命令草案都应由用户结合现场版本、当前配置和变更窗口审阅。", caution=True)

    add_heading(doc, "2. 启动前准备", 1)
    add_table(doc, ["项目", "要求或建议"], [
        ["发布目录", "必须保留完整的 release\\NetworkAutomation\\ 目录，包括 NetworkAutomation.exe 与 _internal。不能只复制 exe 文件。"],
        ["操作系统", "Windows 10/11 x64。首次使用需要 Microsoft Edge WebView2 Runtime，通常已随系统安装。"],
        ["CHM 导入", "安装 7-Zip 并保证 7z.exe 可被系统找到。CHM 使用 7-Zip 解包；不要依赖 hh.exe 的静默解包结果。"],
        ["模型服务", "LLM 与 Embedding 均使用 OpenAI 兼容接口。可先只使用 LLM；未建 Embedding 时会退化到命令名和全文检索。"],
        ["现场设备", "仅对获得授权的交换机填写 SSH 资料和执行写入。准备变更窗口、现网备份与设备侧验收方式。"],
    ], [2100, 7260])

    add_heading(doc, "3. 启动、关闭与本地数据", 1)
    add_heading(doc, "3.1 启动桌面版", 2)
    add_list(doc, [
        "进入 release\\NetworkAutomation\\ 文件夹。",
        "双击 NetworkAutomation.exe。首次启动会初始化本机服务和前端资源。",
        "程序会打开桌面窗口；服务仅监听 127.0.0.1，不向局域网开放。",
    ], numbered=True)
    add_heading(doc, "3.2 开发模式", 2)
    add_body(doc, "需要修改源码时，在项目根目录分别启动后端和前端：")
    add_code(doc, "uv sync --extra dev\nuv run uvicorn app.main:app --reload --app-dir apps/api\n\npnpm install --frozen-lockfile\npnpm --filter network-automation-web dev")
    add_heading(doc, "3.3 本地数据与关闭", 2)
    add_body(doc, "桌面版运行数据位于 %LOCALAPPDATA%\\NetworkAutomation\\data，包含 SQLite 数据库、导入的手册、抽取结果、向量、导出文件、执行记录与日志。替换 release 目录不会删除它。正常关闭窗口即可；正在运行的手册抽取、Embedding 构建或规划任务会在下次启动后显示状态。")

    add_heading(doc, "4. 首次体验：完成一次三交换机组网", 1)
    add_body(doc, "下面用一个不涉及真实下发的示例走完整流程。该示例验证“拓扑 + 需求 -> 每台设备命令草案”，不应直接用于生产网络。")
    add_table(doc, ["对象", "连接与目标"], [
        ["SW1", "连接 PC1、PC2 与 SW3。PC1 使用 VLAN 10，PC2 使用 VLAN 20。"],
        ["SW2", "连接 PC3、PC4 与 SW3。PC3 使用 VLAN 10，PC4 使用 VLAN 20。"],
        ["SW3", "连接 SW1 与 SW2，承载两条上联并提供 VLAN 间三层互通所需的角色。"],
        ["业务需求", "PC1 与 PC3 属于 VLAN 10；PC2 与 PC4 属于 VLAN 20；两个 VLAN 之间允许通信。地址和网关由规划思路补充并由用户审阅。"],
    ], [2100, 7260])
    add_list(doc, [
        "在“设置”填写 LLM 连接信息并保存；若要体验向量检索，再填写 Embedding 信息。",
        "在“手册管理”导入一份已授权的华为 VRP 命令参考手册，等待状态为 completed 或 completed_with_issues；如已配置 Embedding，可为该手册点击“构建 Embedding”。",
        "在“拓扑编辑”添加 SW1、SW2、SW3 和四台 PC。使用“连线（绳）”依次点击两个设备，形成每一条直线。点击每条线，在右侧填写两端接口，例如 SW1 的 GE0/0/1 与 PC1 的 Ethernet0/0/1。保存拓扑。",
        "在“配置规划”选择刚保存的拓扑和该手册，输入业务需求，点击“第一步：生成配置思路”。",
        "检查模型给出的设备角色、VLAN、上联、网关与实施顺序。可以直接修改；思路非空后点击“确认思路并生成命令”。",
        "等待右侧节点状态完成。依次查看 SW1、SW2、SW3 的命令卡片和每页 5 条的手册证据；必要时直接改命令并保存。",
        "仅为体验可到“下发与结果”查看已保存命令，不输入 SSH 密码、不提交下发。若在授权实验环境下执行，仍应逐台填写 SSH 信息并确认。",
    ], numbered=True)

    add_heading(doc, "5. 设置：LLM 与 Embedding", 1)
    add_heading(doc, "5.1 LLM 设置", 2)
    add_table(doc, ["字段", "填写方式"], [
        ["Base URL", "填写模型提供方的 OpenAI 兼容地址，例如 https://<HOST>/v1/。"],
        ["API Key", "填写实际密钥；手册、截图和文档中只应使用占位符。保存后留空表示保持本机已保存的密钥不变。"],
        ["Model", "填写服务端公开的模型名。"],
        ["Temperature", "建议从 0 到 0.3 开始；提高数值会增加表达差异。"],
        ["thinking 策略", "自适应：只在需求理解与命令规划等推理节点启用；始终开启：所有 LLM 节点请求推理；始终关闭：用于兼容或排障。"],
    ], [2100, 7260])
    add_body(doc, "点击“保存本机设置”，再点击“测试 LLM 连接”。请求固定使用 httpx verify=False，并且会兼容不同供应商对 chat_template_kwargs.enable_thinking 与 thinking.type 的支持差异。只应连接可信的内网或私有端点。")
    add_heading(doc, "5.2 Embedding 设置", 2)
    add_table(doc, ["字段", "填写方式"], [
        ["Base URL", "可填写 /v1/ 基地址，也可填写完整的 /v1/embeddings 地址，程序会规范化。"],
        ["Model", "填写 Embedding 模型名，例如 <EMBEDDING_MODEL>。"],
        ["向量维度", "服务要求显式维度时填写，例如 Qwen3-Embedding-4B 可填写 2560；不要求时留空。"],
        ["每批请求条数", "默认 2，可设为 1 至 20。端点限流或并发能力弱时用较小数值。"],
    ], [2100, 7260])
    add_body(doc, "Embedding 只用于把手册页面和检索查询转成向量，参与混合检索排序；它不生成命令。向量保存在本机 SQLite，运行时用 CPU 计算相似度。")

    add_heading(doc, "6. 手册管理：导入、抽取与向量索引", 1)
    add_heading(doc, "6.1 导入与抽取", 2)
    add_list(doc, [
        "进入“手册管理”，选择导入文件。支持 CHM、PDF、HTML/HTM、TXT 与 Markdown。品牌、版本等字段默认可不填；导入后也可以编辑名称和元数据。",
        "观察任务状态、页数、命令数、型号辅助信息和失败项。completed 表示完成；completed_with_issues 表示主体完成但存在失败页；failed 表示需查看具体错误。",
        "CHM 会先由 7-Zip 解包为 HTML，再读取目录、章节与命令页。PDF、HTML 与文本则各自走对应的文本和结构抽取分支。",
        "完成抽取后，可点击“构建 Embedding”。构建过程会按设置页的批大小发送手册页面文本；未构建不阻止后续规划。",
    ], numbered=True)
    add_heading(doc, "6.2 手册中保存了什么", 2)
    add_body(doc, "系统保留原始页路径、标题、命令名称、语法、命令视图、参数、适用范围、前置条件、限制、注意事项与示例片段。它同时保留原始文本页，用于全文与向量检索。型号、系列与版本是抽取辅助信息，不再要求在拓扑中选择型号，也不会成为生成命令的门禁。")
    add_heading(doc, "6.3 单本导入、导出、编辑与删除", 2)
    add_list(doc, [
        "编辑可修改显示名称、品牌、版本和 CLI 方言；未知信息保持为空或“未标注”，不要填入猜测值。",
        "导出会打包该手册的原文件、抽取内容、命令知识、索引元数据和已构建向量。桌面版可选择保存位置，完成后会提示实际路径。",
        "导入单本手册归档时，若发现同名项，系统会询问是否覆盖。",
        "删除只删除当前手册及其本机抽取资料；已存在的历史任务不会被自动改写。",
    ], numbered=True)

    add_heading(doc, "7. 拓扑编辑：设备、直线链路与接口", 1)
    add_heading(doc, "7.1 建立节点", 2)
    add_body(doc, "点击“交换机”或“PC 终端”添加设备，拖动设备调整位置。每个设备的数据互不继承：IP、掩码前缀、网关均可为空；交换机额外可填 SSH IP、端口、用户名和受保护端口。PC 没有 SSH 表单，也不会执行 PC ping。")
    add_body(doc, "交换机可点击“仅查询型号/版本”，它只执行 display version，所得型号和版本仅作审计显示；任务选择手册时不依赖型号映射。")
    add_heading(doc, "7.2 建立与编辑链路", 2)
    add_list(doc, [
        "点击“连线（绳）”。先点第一台设备，再点第二台设备，得到一根从两个设备中心连出的直线。节点上不会显示端点圆点。",
        "点击直线，在右侧填写两端真实接口名并保存，例如 GE0/0/1 与 Ethernet0/0/1。接口标签显示在直线两端附近。",
        "端口书写会原样带入命令。填写 GE0/0/0 时不会自动替换为 GigabitEthernet0/0/0；系统仅在去重、受保护端口比对时将二者视为同一接口。",
        "右键节点或直线可删除。删除节点时，会同时删除连接到它的链路。",
    ], numbered=True)
    add_heading(doc, "7.3 保存、打开与迁移", 2)
    add_list(doc, [
        "填写拓扑名称并点击保存。之后的保存更新当前拓扑；改名会同步更新已保存拓扑的名称。",
        "在“已保存拓扑”选择一项即可打开对应图。可删除当前保存拓扑；存在引用该拓扑的任务时，系统会拒绝删除。",
        "可导出单个拓扑及其当前配置要求；导入同名归档时，界面会询问是否覆盖。",
        "点击“前往配置规划”可带着当前已保存拓扑进入下一页。",
    ], numbered=True)

    add_heading(doc, "8. 配置规划：思路、检索与逐设备命令", 1)
    add_heading(doc, "8.1 两阶段操作", 2)
    add_list(doc, [
        "选择已保存拓扑和一份已完成抽取的手册，填写原始需求。一次任务只使用选择的这一本手册。",
        "点击“第一步：生成配置思路”。模型收到完整拓扑、每台设备信息、IP/掩码/网关（缺失字段明确写为未提供）、真实链路和两端接口，以及原始需求。",
        "审阅模型的思路。它既可提出实施方案，也可补充需求中缺少的地址、网关、角色、顺序或约束。用户可任意改写。",
        "思路非空后，点击“确认思路并生成命令”。此操作会清空旧命令和旧进度，但保留最终思路、拓扑、手册和原始需求。",
        "生成完成后，设备命令从左到右排列。点击一个设备查看并修改该设备命令，使用“保存命令”写入此设备的最终版本。",
    ], numbered=True)
    add_heading(doc, "8.2 命令如何从手册中找到", 2)
    add_body(doc, "命令阶段由 LangGraph 在应用内部编排，不假设模型具备原生 function calling。模型输出结构化检索意图，程序执行检索并将结果带回下一节点。流程最多两轮：每轮模型可给出最多 5 个重要查询词；程序用命令名精确匹配、SQLite FTS5 全文检索和可选 Embedding 余弦相似度做混合召回。")
    add_body(doc, "每一轮的候选按手册页面去重。模型从候选中挑选当前任务真正需要的页面，判断证据是否足够；不足则进行第二轮。两轮结束后，即使部分命令未找到，系统仍会让模型基于已找到页面和当前拓扑写出可编辑草案，而不是清空命令面板。")
    add_heading(doc, "8.3 右侧进度、停止与重新开始", 2)
    add_list(doc, [
        "右侧栏实时显示任务走到的节点及状态，例如任务初始化、手册检索、页面选择、某台设备的命令规划与完成。页面切换后工作状态不会消失。",
        "右侧不展示模型的内部思考文本、JSON 或正式输出，避免干扰审阅。正式结果体现在配置思路、证据和设备命令中。",
        "点击“停止”会请求取消当前规划。点击“重新开始并生成配置思路”会先停止旧任务，并清空思路、命令和进度后重新起草。",
        "命令生成期间又点击确认时，系统会继续订阅已在运行的任务，避免必须第二次点击才出现进度。",
    ], numbered=True)
    add_callout(doc, "审阅重点", "检查设备角色、真实接口、厂商配置视图、地址与网关、上联承载、保护端口和验证命令。命令草案允许自由修改；应用不会替代变更评审。", caution=True)

    add_heading(doc, "9. 模板管理：保存、查看与迁移", 1)
    add_body(doc, "在配置规划中已有设备命令后，可以点击“保存为模板”，填写标题和简介。模板只保存当前拓扑、需求、配置思路与各设备的最终命令，不保存模型思考、JSON、密码或内部审计文本。")
    add_list(doc, [
        "在“模板管理”查看、编辑标题和简介、导出、导入或删除单个模板。",
        "模板详情的拓扑图与拓扑编辑器一致：可拖动设备查看，并显示从设备中心出发的直线以及两端接口名称。模板预览仅供浏览，不提供连线编辑。",
        "导入同名模板会询问是否覆盖；导出时可选择保存位置，成功后界面会提示实际路径。",
        "模板不会自动套用到新任务，也不会作为 LLM 提示词。它是人工可查阅、可迁移的历史快照，避免旧设备命令误导新拓扑。",
    ], numbered=True)

    add_heading(doc, "10. 下发与结果：逐台执行与设备回显", 1)
    add_heading(doc, "10.1 下发前编辑和保存", 2)
    add_list(doc, [
        "选择任务和设备计划，在“本设备待执行命令”中查看或编辑命令。",
        "命令非空时，“保存命令”可点击。编辑后必须保存，保存后的版本才会被下发。即使没有检测到差异，也允许手动再次保存。",
        "填写本次 SSH 地址、端口、用户名和密码。密码只用于这一次连接，不写入 SQLite、模板或导出文件。",
        "确认“仅下发 <设备名>”。每次只针对当前设备，应用不会批量发送。",
    ], numbered=True)
    add_heading(doc, "10.2 执行行为与回显", 2)
    add_body(doc, "程序通过 Netmiko 的 send_config_set 一次提交当前设备完整配置块，再以本地 SSE 和轮询在右侧“设备实时回显”中显示连接、配置、验证和保存的完整设备输出。设备名和接口上下文以设备返回的提示符或命令回显为准。")
    add_body(doc, "独立一行的第一个 return 被视为本次配置块结束；其后的内容不会发送，并应在审阅阶段删除或移动。执行完成后，下方记录会保留状态、预检、错误信息、save 结果和完整回显。验证失败时不自动 save；验证成功才自动执行保存。")
    add_callout(doc, "受保护端口", "在拓扑中标记的受保护端口不会被自动生成或写入。该保护只依据当前拓扑和当前任务；不要把一次实验的临时限制误解为永久设备策略。", caution=True)

    add_heading(doc, "11. 导入、导出、删除与备份", 1)
    add_table(doc, ["对象", "可操作内容", "注意事项"], [
        ["拓扑", "保存、打开、导出、导入、删除。", "归档包含图、节点、链路接口和当前配置要求；同名导入会询问覆盖。"],
        ["手册", "编辑、构建 Embedding、导出、导入、删除。", "归档包含原始资料、抽取内容和索引。确认厂商版权和内部资料边界后再迁移。"],
        ["模板", "查看、编辑标题/简介、导出、导入、删除。", "只含可审阅的拓扑、需求、思路和命令，不含密码。"],
        ["本地 data", "系统自动维护。", "不要直接复制正在使用的 SQLite 文件作为迁移方式；优先使用页面的单项导入/导出。"],
    ], [1700, 3400, 4260])

    add_heading(doc, "12. 常见问题与排查", 1)
    add_table(doc, ["现象", "先检查什么", "处理建议"], [
        ["导入手册没有开始解析", "状态、文件格式、7-Zip、日志。", "刷新手册管理页查看任务；CHM 确认 7z.exe 可用。失败时查看失败项后重新导入。"],
        ["LLM 连接失败", "Base URL、模型名、API Key、内网连通性。", "先保存，再点击测试 LLM 连接；用 thinking=off 判断是否为推理扩展兼容问题。"],
        ["Embedding 构建失败", "URL、模型、dimensions、批大小。", "将每批请求条数降为 1 或 2；确认端点接受 input 数组与可选 dimensions。"],
        ["保存拓扑失败", "拓扑名称、每个节点字段、浏览器错误提示。", "确保已填写拓扑名称；刷新后重新打开已保存拓扑，不要让未保存页面长时间停留。"],
        ["命令生成没有进度", "右侧节点状态、任务是否仍在执行、LLM 连通性。", "先等待节点推进；重复点击确认会继续订阅任务。需要重做时用停止或重新开始。"],
        ["下发页保存命令不可用", "命令是否非空、是否正在执行。", "结束当前执行或选择另一设备后再保存。非空命令即使未改动也可手动保存。"],
        ["EXE 打开无响应", "是否只复制了 exe、_internal 是否存在、WebView2、日志。", "从完整 release\\NetworkAutomation\\ 目录启动；检查 %LOCALAPPDATA%\\NetworkAutomation\\data\\logs。"],
    ], [1850, 3000, 4510])

    add_heading(doc, "13. GitHub、许可证与二次开发", 1)
    add_body(doc, "项目代码可从 https://github.com/chenjinyu0723/network-automation 获取。仓库采用 MIT License，欢迎在理解厂商手册授权、组织安全要求和现场变更流程的前提下继续优化手册抽取、检索、规划、设备适配与测试。")
    add_body(doc, "提交代码时不要上传本机 data、数据库、API Key、SSH 密码、导入的厂商手册、设备回显或构建产物。仓库的 .gitignore 已覆盖这些常见本地文件。")
    add_code(doc, "git clone https://github.com/chenjinyu0723/network-automation.git\ncd network-automation\nuv sync --extra dev\npnpm install --frozen-lockfile\n\n# 重新构建桌面版\nuv sync --extra desktop\n.\\scripts\\build_desktop.ps1")

    add_heading(doc, "附录：操作前检查清单", 1)
    add_list(doc, [
        "已从完整 release\\NetworkAutomation\\ 目录启动应用，且后端状态正常。",
        "LLM 已保存并通过连通性测试；如果使用向量检索，Embedding 已配置且目标手册已构建索引。",
        "所选手册已完成抽取，且与目标设备厂商、命令风格和现场版本尽量一致。",
        "拓扑已保存；所有需要配置的真实链路和两端接口都已填写。未提供的 IP、掩码、网关保持为空，而不是继承其它设备的值。",
        "配置思路已由人审阅并补全；每台设备命令已查看、必要时修改并保存。",
        "确认 SSH 目标、账号权限、受保护端口、变更窗口、现网备份和设备侧验证方式。",
        "准备逐台确认下发；不把命令草案、模型输出或模板快照当作生产变更的最终批准。",
    ], numbered=False)

    doc.core_properties.title = "AI Agent 工业交换机自动配置用户手册"
    doc.core_properties.subject = "当前桌面版功能与操作说明"
    doc.core_properties.author = "network-automation"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_manual()
