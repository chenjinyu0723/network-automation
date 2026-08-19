"""Build the end-user manual for the local NetworkAutomation desktop app.

The file is intentionally kept in the repository so future updates to the
application can regenerate the DOCX from one maintained source.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "工业交换机自动配置用户手册.docx"
FONT = "Microsoft YaHei"
MONO_FONT = "Cascadia Mono"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17365D"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
CAUTION = "FFF6E5"
WHITE = "FFFFFF"
CONTENT_DXA = 9360


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid, width in zip(grid_cols, widths):
        grid.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_external_hyperlink(paragraph, url, label=None):
    label = label or url
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    r_pr.append(fonts)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = label
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("AI Agent 工业交换机自动配置 | 本地单用户版")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("NetworkAutomation  |  第 ")
    set_run_font(r, size=8.5, color=MUTED)
    add_page_field(footer)
    r = footer.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)


def add_para(doc, text="", bold_prefix=None, align=None, after=None, before=None, color=None, size=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, size=size, color=color, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run, size=size, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def _new_numbering_id(doc):
    """Clone the List Number abstract definition and restart it at 1."""
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]._element
    style_ppr = style.find(qn("w:pPr"))
    style_numpr = style_ppr.find(qn("w:numPr")) if style_ppr is not None else None
    style_numid = style_numpr.find(qn("w:numId")) if style_numpr is not None else None
    base_numid = style_numid.get(qn("w:val")) if style_numid is not None else None
    base_num = next(
        (node for node in numbering.findall(qn("w:num")) if node.get(qn("w:numId")) == base_numid),
        None,
    )
    if base_num is None:
        raise RuntimeError("List Number numbering definition is unavailable")
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    max_numid = max((int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))), default=0)
    new_numid = str(max_numid + 1)
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), new_numid)
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    new_num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    new_num.append(override)
    numbering.append(new_num)
    return new_numid


def add_number(doc, text):
    previous_is_number = bool(doc.paragraphs) and doc.paragraphs[-1].style.name == "List Number"
    if not previous_is_number:
        doc._manual_numbering_id = _new_numbering_id(doc)
    p = doc.add_paragraph(style="List Number")
    p_pr = p._p.get_or_add_pPr()
    existing_numpr = p_pr.find(qn("w:numPr"))
    if existing_numpr is not None:
        p_pr.remove(existing_numpr)
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), doc._manual_numbering_id)
    numpr.append(ilvl)
    numpr.append(numid)
    p_pr.append(numpr)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_note(doc, title, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    set_run_font(r, size=10, color=INK, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10, color="333333")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name=MONO_FONT, size=9, color="2F3A4A")
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, value, bold=True, color=INK, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_h1(doc, title):
    return doc.add_heading(title, level=1)


def add_h2(doc, title):
    return doc.add_heading(title, level=2)


def add_h3(doc, title):
    return doc.add_heading(title, level=3)


def add_toc(doc, toc_pages):
    add_h1(doc, "目录")
    add_para(doc, "本手册按完整操作顺序组织。Word 的“导航窗格”可展开并定位二级标题；下列页码为一级章节定位。", color=MUTED, size=9.5)
    entries = [
        ("1. 认识应用", "1"), ("2. 启动前准备", "2"), ("3. 启动应用", "3"),
        ("4. 从零完成一次组网", "4"), ("5. 设置模型服务", "5"), ("6. 手册管理", "6"),
        ("7. 拓扑编辑", "7"), ("8. 配置规划", "8"), ("9. 下发、验证与撤销", "9"),
        ("10. 模板管理", "10"), ("11. 单项导入、导出与删除", "11"), ("12. 常见问题", "12"),
        ("13. 本地数据与安全", "13"), ("14. GitHub 与二次开发", "14"), ("附录 A. Git 忽略规则", "A"),
        ("附录 B. 上线前检查清单", "B"),
    ]
    left, right = entries[:8], entries[8:]
    table = doc.add_table(rows=max(len(left), len(right)), cols=2)
    set_table_geometry(table, [4680, 4680])
    for row_index in range(len(table.rows)):
        for col_index, source in enumerate((left, right)):
            cell = table.cell(row_index, col_index)
            if row_index < len(source):
                label, key = source[row_index]
                page = toc_pages.get(key, "-")
                page_label = "-" if page == "-" else f"第 {page} 页"
                set_cell_text(cell, f"{label}  ·  {page_label}", color=INK, size=9.5)
            else:
                set_cell_text(cell, "", size=9.5)
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                element = borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "nil")
    doc.add_page_break()


def build_document(toc_pages):
    doc = Document()
    configure_document(doc)

    # Cover
    for _ in range(6):
        add_para(doc, "", after=0)
    p = add_para(doc, "本地单用户版", align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, size=12, after=8)
    p.runs[0].bold = True
    p = add_para(doc, "AI Agent 工业交换机\n自动配置用户手册", align=WD_ALIGN_PARAGRAPH.CENTER, color=INK, size=25, after=12)
    p.runs[0].bold = True
    add_para(doc, "从手册注入、手绘拓扑到逐台确认下发", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=13, after=28)
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [2700, 6660])
    for row_index, label, value in (
        (0, "适用范围", "Windows 桌面应用 / 本地 FastAPI / React 工作台"),
        (1, "更新日期", "2026-08-18"),
        (2, "项目地址", "https://github.com/chenjinyu0723/network-automation"),
    ):
        label_cell = table.cell(row_index, 0)
        set_cell_shading(label_cell, LIGHT_BLUE)
        set_cell_text(label_cell, label, bold=True, color=INK, size=10)
        set_cell_text(table.cell(row_index, 1), value, size=10)
    add_para(doc, "本手册讲解功能使用与审阅流程。设备命令是否适用于现场，始终应由具备网络变更权限的人员结合手册、现网和变更流程确认。", align=WD_ALIGN_PARAGRAPH.CENTER, color=MUTED, size=9.5, before=26)
    doc.add_page_break()

    add_toc(doc, toc_pages)

    add_h1(doc, "1. 认识应用")
    add_para(doc, "AI Agent 工业交换机自动配置是一个在本机运行的单用户应用。它把用户导入的交换机手册抽取为命令知识库，结合手绘拓扑、配置需求与可选模板，生成按设备拆分的配置命令草案。应用的目标是把“拓扑 + 需求”推进到“每台设备都可审阅的命令”，而不是绕过人工直接改网。")
    add_note(doc, "核心原则", "系统可以自动分析、检索、规划和生成草案；命令下发必须由用户逐台确认。模型输出、手册证据、拓扑推断和执行回显会保留在任务中，方便复核。")
    add_h2(doc, "1.1 完整工作流")
    add_number(doc, "在“设置”页配置 LLM 和可选 Embedding 服务，并确认 LLM 连接。")
    add_number(doc, "在“手册管理”页导入一份或多份手册；完成抽取后，按需构建 Embedding 索引。")
    add_number(doc, "在“拓扑编辑”页放置交换机和 PC，用连线工具连接设备，并填写每条线两端的真实接口名。")
    add_number(doc, "保存拓扑后，在“配置规划”页选择拓扑、手册，输入需求，先生成并确认配置思路。")
    add_number(doc, "配置思路非空后，生成逐设备命令草案，查看手册证据、LLM 输出和静态提示；必要时编辑命令。")
    add_number(doc, "在“下发与结果”页逐台确认发送，实时查看回显和验证结果。成功验证后才会自动 save；可选执行 PC ping 验收。")
    add_h2(doc, "1.2 当前边界")
    add_bullet(doc, "每个配置任务由用户选择一份已完成抽取的手册作为本轮唯一命令上下文；无需在拓扑中填写型号。")
    add_bullet(doc, "PC 图标默认不支持 SSH。只有在用户填写并授权 PC SSH 信息时，才可在验证阶段运行受限的 ping 验收。")
    add_bullet(doc, "未知厂商或不常见需求不会因为缺少“内置能力”被直接拦截：系统会走通用手册检索和命令草案路径，但会明显提示证据不足或未验证状态。")
    add_bullet(doc, "系统不会自动批量下发，也不代替变更审批。保护端口不会进入自动生成和写入范围。")

    add_h1(doc, "2. 启动前准备")
    add_h2(doc, "2.1 桌面版运行条件")
    add_table(doc, ["项目", "用途与检查"], [
        ("发布目录", "保留整个 release\\NetworkAutomation\\ 目录。NetworkAutomation.exe 依赖同目录的 _internal，不能只复制一个 EXE。"),
        ("Windows", "Windows 10/11 x64。首次启动需要 Microsoft Edge WebView2 Runtime；多数 Windows 10/11 已自带。"),
        ("CHM 导入", "安装 7-Zip，并确保 7z.exe 可从系统 PATH 找到。CHM 应由 7-Zip 解包，不能依赖 hh.exe。"),
        ("模型服务", "LLM 和 Embedding 使用 OpenAI 兼容接口。没有模型也可查看手册、画拓扑和使用关键词/FTS 检索；生成质量会下降。"),
        ("网络设备", "仅在有明确授权的实验或生产窗口中配置 SSH 信息。不要把未授权设备或保护端口放入可写范围。"),
    ], [2200, 7160])
    add_h2(doc, "2.2 本机数据位置")
    add_para(doc, "桌面版的运行数据固定保存在 %LOCALAPPDATA%\\NetworkAutomation\\data。这里包含 SQLite 数据库、导入手册原件、解包与抽取结果、Embedding 向量、单项导出文件、日志和执行审计。删除或替换 release 目录不会自动删除这些数据。")
    add_note(doc, "数据迁移提示", "需要转移到另一台电脑时，优先使用页面中的“单项导出/导入”。直接复制 data 目录可能会携带未导出的凭据引用和环境路径，不建议作为常规迁移方式。", CAUTION)

    add_h1(doc, "3. 启动应用")
    add_h2(doc, "3.1 推荐：双击桌面版")
    add_number(doc, "打开 release\\NetworkAutomation\\。")
    add_number(doc, "双击 NetworkAutomation.exe，第一次启动请等待本机后端和网页资源初始化。")
    add_number(doc, "应用会打开原生桌面窗口，内部服务只绑定 127.0.0.1，不监听局域网地址。看到左侧导航栏即表示启动完成。")
    add_note(doc, "窗口无响应时", "先完全退出正在运行的 NetworkAutomation.exe，再从完整发布目录重新启动。仍无法打开时检查 WebView2 Runtime、杀毒软件拦截和 %LOCALAPPDATA%\\NetworkAutomation\\data\\logs 中的日志。", CAUTION)
    add_h2(doc, "3.2 开发模式启动")
    add_para(doc, "二次开发或排障时，可在项目根目录打开两个 PowerShell 窗口，分别启动后端与前端：")
    add_code(doc, [
        "# 窗口 A：后端",
        "uv sync --extra dev",
        "uv run uvicorn app.main:app --reload --app-dir apps/api",
        "",
        "# 窗口 B：前端",
        "pnpm install --frozen-lockfile",
        "pnpm --filter network-automation-web dev",
    ])
    add_para(doc, "开发模式下后端默认地址为 http://127.0.0.1:8000，前端默认地址为 http://127.0.0.1:5173。请从项目根目录执行 pnpm 安装；不要在 apps/web 子目录单独执行 pnpm install。")
    add_h2(doc, "3.3 关闭应用")
    add_para(doc, "关闭桌面窗口会结束本地服务。手册抽取、Embedding 构建和规划任务的状态会写入 SQLite；下次启动后可在对应页面查看结果或按提示重试。正在流式生成的规划任务可以先在“配置规划”右侧点击“停止”。")

    add_h1(doc, "4. 从零完成一次组网")
    add_para(doc, "下面的体验路径使用一个不会自动下发的最小案例。它覆盖从模型设置到生成命令的完整流程；若要实际发送，请只在获得授权的隔离设备上继续第 9 章。")
    add_h2(doc, "4.1 体验案例")
    add_table(doc, ["对象", "本次填写内容"], [
        ("拓扑", "PC1 -- SW1 -- PC2。PC1 与 SW1 的真实接口填 Ethernet0/0/1 和 GE0/0/1；PC2 与 SW1 的真实接口填 Ethernet0/0/2 和 GE0/0/2。"),
        ("手册", "选择已完成抽取的华为 S1700/S5700/S6700 V600R025C00 命令参考，或选择你的目标厂商手册。"),
        ("需求", "在 SW1 上创建 VLAN 10，将连接 PC1 和 PC2 的两个端口配置为 Access 并加入 VLAN 10。"),
        ("预期", "先得到可编辑的实施思路；确认后得到按 SW1 拆分的命令草案和对应手册证据。命令要由用户自行审阅。"),
    ], [1800, 7560])
    add_h2(doc, "4.2 按顺序操作")
    add_number(doc, "打开“设置”，填写 LLM 信息并点击“保存本机设置”和“测试 LLM 连接”。Embedding 可稍后再配置。")
    add_number(doc, "打开“手册管理”，导入手册并等待任务状态成为 completed 或 completed_with_issues。已配置 Embedding 时，点击“构建 Embedding”。")
    add_number(doc, "打开“拓扑编辑”，添加 SW1、PC1、PC2。选择连线工具，依次单击两个设备建立直线；单击线，在右侧填写两端的真实接口。")
    add_number(doc, "点击“保存拓扑”，给出清晰名称，例如“VLAN10 体验”。保存成功后，点击“前往配置规划”。")
    add_number(doc, "在“配置规划”选择已保存拓扑和目标手册，输入上面的需求。可不选模板。点击“第一步：生成配置思路”。")
    add_number(doc, "阅读并修改配置思路。思路必须明确 VLAN、涉及设备、真实接口、实施顺序和限制；确认无误后点击“保存配置思路”。")
    add_number(doc, "点击“第二步：生成配置命令”。在设备卡片查看 SW1 命令、手册命令证据与右侧任务进度。没有确认过的思路时，这个按钮不可用。")
    add_number(doc, "确认命令符合手册和现场后，可在“下发与结果”逐台下发；只想体验规划时到此结束即可。")
    add_note(doc, "体验重点", "真实接口文本会原样进入命令。填写 GE0/0/1 时，应用会生成 interface GE0/0/1；填写 GigabitEthernet0/0/1 时也会原样保留。系统只在端口去重、保护端口和回显比对时把两种写法视为同一物理口。")

    add_h1(doc, "5. 设置模型服务")
    add_para(doc, "进入左侧“设置”。所有模型调用采用 OpenAI 兼容接口，应用的 HTTP 传输固定使用 verify=False，因此只应配置可信的内网或私有服务。页面保存后不会回显 API Key；密钥只保存在本机系统凭据存储。")
    add_h2(doc, "5.1 LLM 设置")
    add_table(doc, ["字段", "怎么填", "作用"], [
        ("Base URL", "填写供应商提供的 OpenAI 兼容服务基地址。", "用于聊天补全请求。"),
        ("API Key", "填写实际密钥；手册和截图中请使用占位符。", "请求鉴权，不写入 SQLite。"),
        ("Model", "填写服务端发布的聊天模型名称。", "决定规划、命令草案和审阅使用的模型。"),
        ("Temperature", "建议从 0 到 0.3 开始。", "数值低时输出更稳定；需求探索可适度提高。"),
        ("thinking 策略", "选择 adaptive、always 或 off。", "控制是否向兼容端点请求思考模式。"),
    ], [1550, 4300, 3510])
    add_bullet(doc, "adaptive（推荐）：仅在需求理解、命令计划/修订、命令审阅和结果诊断等推理节点开启；检索、静态校验、执行和 save 不调用 thinking。")
    add_bullet(doc, "always：全部实际 LLM 调用请求 thinking，适合测试推理质量，耗时和消耗可能更高。")
    add_bullet(doc, "off：所有 LLM 调用关闭 thinking，适合排查不支持扩展字段的端点。")
    add_para(doc, "应用会同时兼容 chat_template_kwargs.enable_thinking 与 thinking.type；若端点拒绝其中一种扩展字段，会自动按单字段、再无扩展字段的顺序重试，并在任务审计中记录降级原因。")
    add_h2(doc, "5.2 Embedding 设置")
    add_para(doc, "Embedding 用于把手册命令、参数、说明和示例转成向量，并在规划阶段辅助寻找语义接近的命令。它不是命令生成器，也不是唯一检索通道。未配置 Embedding 时，系统会回退到命令名精确匹配和 SQLite FTS5 全文检索。")
    add_table(doc, ["字段", "填写建议"], [
        ("Embedding Base URL", "可填写 /v1/ 基地址，或完整的 /v1/embeddings 地址。"),
        ("Model", "填写供应商公布的向量模型名，例如 Qwen3-Embedding-4B。"),
        ("Dimensions", "服务要求显式维度时填写，如 2560；留空则不发送 dimensions 字段。"),
        ("默认批量", "默认 2，可在 1-20 间调整。端点限流严格或并发能力弱时使用较小值。"),
    ], [2400, 6960])
    add_note(doc, "公司接口示例", "若服务提供 POST http://adn-ai.test.huawei.com:5000/v1/embeddings，可将完整地址填入 Embedding Base URL，模型填 Qwen3-Embedding-4B，维度填 2560。API Key 必须从你的授权渠道获取，文档中不要记录真实值。")
    add_h2(doc, "5.3 连通性检查")
    add_number(doc, "填写 LLM 的 Base URL、API Key、Model 和 Temperature 后，点击“保存本机设置”。")
    add_number(doc, "点击“测试 LLM 连接”。成功后可继续规划；失败时依次核对 URL、模型名、Key、公司网络和端点是否支持 Chat Completions。")
    add_number(doc, "Embedding 的实际可用性在“构建 Embedding”后确认。索引任务会显示成功、失败原因和进度。")

    add_h1(doc, "6. 手册管理")
    add_para(doc, "手册管理的目标是把任意品牌、版本与格式的命令文档变成同一种可检索知识结构。配置任务不再要求从型号库挑设备型号，只需选择与目标设备匹配的已完成手册。")
    add_h2(doc, "6.1 导入一份手册")
    add_number(doc, "进入“手册管理”，把 CHM、PDF、HTML、HTM、TXT 或 Markdown 文件拖入导入区，或点击选择文件。")
    add_number(doc, "填写品牌、版本、名称等可编辑信息。未知时保持“无”，后续可编辑或重命名。")
    add_number(doc, "观察“当前导入任务”的阶段、页数、命令数和详细信息。完成后，这份手册会出现在任务选择列表。")
    add_h2(doc, "6.2 各格式如何进入知识库")
    add_table(doc, ["输入格式", "处理方式", "使用注意"], [
        ("CHM", "由本机 7-Zip 解包为 HTML，再解析目录、章节与命令页。", "安装 7-Zip；不要使用 hh.exe 的解包结果作为依据。"),
        ("PDF", "读取页面文字、标题和章节结构，再按页面/章节抽取。", "扫描版或复杂双栏版需要关注文字提取质量。"),
        ("HTML / HTM", "保留目录层级，抽取正文、命令格式、参数和示例。", "尽量导入完整目录而非单个孤立页面。"),
        ("TXT / Markdown", "按标题、空行和代码块切分，保留原始段落。", "建议写清适用产品、版本和命令上下文。"),
    ], [1700, 4350, 3310])
    add_h2(doc, "6.3 抽取结果与任务状态")
    add_para(doc, "抽取会尽量识别章节标题、命令名称、命令格式、视图、参数、适用范围、约束、注意事项和示例。命令条目会保留来源页与章节，供规划页显示“为什么引用这条命令”。")
    add_table(doc, ["状态", "含义与处理"], [
        ("queued / running", "正在排队或处理。CHM 大手册解包和解析时间较长，请等待页面进度变化。"),
        ("completed", "抽取完成，可用于配置规划。"),
        ("completed_with_issues", "主体完成但部分页面失败。查看失败项；若关键命令缺失，再重试或换来源文件。"),
        ("failed", "任务中断或解析失败。根据详细信息修复 7-Zip、文件可读性或路径问题，然后使用“从断点重试”。"),
    ], [2500, 6860])
    add_h2(doc, "6.4 构建 Embedding 与检索")
    add_number(doc, "确认该手册状态为 completed 或 completed_with_issues，并已在“设置”保存 Embedding 连接。")
    add_number(doc, "点击该手册行的“构建 Embedding”。索引任务会按设置的默认批量把文本发送到配置的向量接口。")
    add_number(doc, "向量以 SQLite BLOB 保存，检索时由 CPU 计算相似度，不依赖云向量数据库。索引完成后，规划会融合命令名、FTS5 和向量结果。")
    add_note(doc, "如何提高命令命中率", "把用户需求写成可检索的目标和约束，例如“在两个交换机互联口允许 VLAN 10、20 通过”比“网络要通”更容易定位正确的 Trunk、VLAN 与允许 VLAN 命令。模型会补充检索词，但仍应由用户审阅证据和最终命令。")
    add_h2(doc, "6.5 编辑、删除、导入和导出")
    add_bullet(doc, "编辑：可修正手册名称、品牌、版本和方言。默认值为“无”，不要把不确定的信息填成假定值。")
    add_bullet(doc, "导出：导出单本手册的原文、抽取结果、命令、型号辅助映射和向量。应用会提示实际保存路径，桌面版可在文件对话框中选择位置。")
    add_bullet(doc, "导入归档：选择单本手册归档后，若发现同名内容，按弹窗确认是否覆盖；覆盖只影响同名手册。")
    add_bullet(doc, "删除：删除当前手册及其本地抽取与向量数据。正在被已有任务引用的历史快照不会被改写，但新任务不能再选择它。")

    add_h1(doc, "7. 拓扑编辑")
    add_para(doc, "拓扑完全由手工绘制，不会通过 LLDP/CDP 自动发现设备。它既是规划输入，也是保护端口和接口范围的依据，所以每条真实链路都应填写两端实际接口。")
    add_h2(doc, "7.1 创建节点")
    add_number(doc, "在设备区选择“交换机”或“PC”，然后单击画布放置节点。交换机可增加上二、下二方向的可用接口槽位；PC 只保留一个接口。")
    add_number(doc, "选中节点，在右侧填写名称。IP、掩码前缀、网关都是每个设备各自可选字段；留空就是未填写，应用不会把一个设备的值当作另一个设备默认值。")
    add_number(doc, "交换机可选填写 SSH IP、端口、账户和密码。PC 默认不支持 SSH；只有配置并授权后才用于 ping 验收。")
    add_h2(doc, "7.2 用“绳子”建立直线")
    add_para(doc, "节点本身不显示小圆点。选择独立的“连线（绳）”工具后，先单击第一台设备，再单击第二台设备，系统会在两个设备方框的正中心生成一条直线。线中间有两个可编辑点，表示两端的实际接口。")
    add_number(doc, "单击一根线，在右侧面板填写 A 端与 B 端接口，例如 GE0/0/1 与 Ethernet0/0/1。")
    add_number(doc, "点击保存。线旁会显示你填写的接口标签，标签只表达物理映射，不替代手册中的端口语法。")
    add_number(doc, "需要删除时，右键单击交换机、PC 或连线，选择删除。删除节点会同时移除关联的线。")
    add_h2(doc, "7.3 接口与保护端口")
    add_table(doc, ["项目", "规则"], [
        ("端口文本", "按输入原样用于命令。GE0/0/0 不会被自动替换成 GigabitEthernet0/0/0，反过来也一样。"),
        ("等价判断", "仅在去重、保护和回显比对时，GE 与 GigabitEthernet 被识别为同一物理接口。"),
        ("保护端口", "在交换机属性中明确标记。系统不会为它生成自动配置，也会在执行前拒绝写入。保护不是永久禁用：用户可在拓扑中修改保护设置后重新规划。"),
        ("未映射端口", "没有出现在链路端口映射中的接口不会被猜测为可写接口。"),
    ], [2100, 7260])
    add_h2(doc, "7.4 保存、打开和单项迁移")
    add_bullet(doc, "点击“保存拓扑”保存当前拓扑。修改后仍是修改当前保存拓扑；改名称会同步更新已保存拓扑的名称，不需要额外维护版本名。")
    add_bullet(doc, "在“已保存拓扑”中选择一个条目即可打开对应画布。打开后可以继续编辑、重新保存，或点击“前往配置规划”。")
    add_bullet(doc, "导出当前拓扑会同时带出当前配置要求/命令快照（如有）且不包含凭据。桌面版会让你选择保存位置，并提示最终路径。")
    add_bullet(doc, "导入单个拓扑归档时，若名称重复，应用会询问是否覆盖；可删除当前保存拓扑，但删除前应确认没有需要保留的任务记录。")

    add_h1(doc, "8. 配置规划")
    add_para(doc, "“配置规划”是两阶段流程：先得到和确认配置思路，再检索手册并生成命令。这样用户可以先核对网络设计，再让模型进入命令层面。思路为空时，应用不会进入命令生成。")
    add_h2(doc, "8.1 创建规划任务")
    add_number(doc, "选择一个已保存拓扑和一份 completed / completed_with_issues 手册。手册选择决定本轮命令上下文。")
    add_number(doc, "输入配置需求。建议包含通信目标、设备角色、VLAN/地址/协议、指定接口、互联关系、不能改动的限制和期望验证方式。")
    add_number(doc, "如有相似的已保存模板，可选择它作为参考。模板仅帮助模型理解角色划分、实施顺序和命令组织，旧设备、端口、VLAN、地址与 CLI 不会被直接套用。")
    add_h2(doc, "8.2 第一步：生成并确认配置思路")
    add_para(doc, "点击“第一步：生成配置思路”。LLM 会根据拓扑和需求给出设备角色、实施顺序、涉及接口和约束。生成后在文本框中审阅；可以直接改写、补充或删减，再点击“保存配置思路”。")
    add_note(doc, "通过标准", "思路至少应说明：哪些设备需要处理、每台设备要实现什么、哪些真实接口参与、先后顺序、互联端口或保护端口的限制、以及怎样验证。没有这些信息时，应先修改思路而不是直接生成命令。", CAUTION)
    add_h2(doc, "8.3 第二步：检索并生成命令")
    add_para(doc, "确认非空思路后点击“第二步：生成配置命令”。LangGraph 会依次运行：读取已确认思路和拓扑事实、LLM 精炼意图与检索词、主动检索手册、LLM 生成受约束 CommandPlan、本地将证据与拓扑端口编译为命令、静态检查，并用独立 LLM 给出审阅提示。模型不需要原生 function calling；工具行为由图节点执行并把结果回填上下文。")
    add_para(doc, "检索采用混合策略：优先命令名精确匹配，其次 SQLite FTS5 全文检索，再叠加可选的 Embedding 语义相似度。检索不到充分证据、模型 JSON 异常或模型完全不可用时，系统仍尽量保留可编辑的 LLM CLI 草案、手册示例或占位参考，但会以“未验证”提示标记，不能伪装成可靠命令。")
    add_h2(doc, "8.4 阅读设备卡片和右侧侧边栏")
    add_table(doc, ["区域", "查看与操作"], [
        ("每台设备卡片", "查看需求、配置思路、检索到的命令证据、生成 CLI、静态提示与 LLM 审阅结果。可以修改命令草案。"),
        ("手册命令证据", "按页显示，默认每页 10 条并可滚动。证据应与目标命令、命令前缀、适用范围和章节一致。"),
        ("规划进度侧栏", "异步流式显示阶段、模型 thinking、正式输出和错误。thinking 完成后会折叠，正式输出保留完整内容。"),
        ("停止", "运行中点击“停止”会写入取消令牌。当前流式响应在可取消位置结束，后续 LangGraph 节点不会继续执行。"),
    ], [2200, 7160])
    add_h2(doc, "8.5 命令审阅要点")
    add_bullet(doc, "检查命令方言与所选手册是否一致。非华为手册不会被自动注入 system-view；未知厂商使用通用手册路径。")
    add_bullet(doc, "检查每个 interface 后面的端口是否来自当前拓扑，且没有碰到保护端口、上联限制或未映射接口。")
    add_bullet(doc, "确认没有把无关的地址、VLAN、策略、save、reboot、reset、delete 等维护动作混入草案。")
    add_bullet(doc, "将黄色“未验证”当作需要人工查手册的信号，而不是“命令正确”的证明。用户可以编辑命令；最终是否发送由用户决定。")

    add_h1(doc, "9. 下发、验证与撤销")
    add_para(doc, "命令生成完成后，进入“下发与结果”。此页使用 Netmiko 通过 SSH 执行单台设备计划；不会提供一键批量下发。密码只在活动连接中使用，不写入任务显示内容。")
    add_h2(doc, "9.1 下发前确认")
    add_number(doc, "选择配置任务和要执行的设备计划。确认命令非空、设备 SSH 信息已在拓扑保存、目标设备和端口在授权范围内。")
    add_number(doc, "再次检查设备卡片的命令和保护端口提示。建议先从影响最小的一台设备开始。")
    add_number(doc, "点击当前设备的确认发送。每台设备都要单独确认，不能自动连续发送。")
    add_h2(doc, "9.2 实时回显、验证和保存")
    add_para(doc, "下发过程中，右侧侧边栏会流式显示 SSH 连接、每条命令、设备回显、验证命令和错误。系统会记录执行前的只读信息（例如 display version）供审计，并对手册证据、审批内容、端口范围和保护端口进行检查。")
    add_table(doc, ["结果", "系统行为"], [
        ("命令和验证成功", "自动执行 save，并记录回显、验证和 save 状态。"),
        ("命令失败或验证失败", "保留错误和设备回显，不自动执行 save。请先分析现场状态与命令，再决定是否修订计划。"),
        ("连接失败", "不会改变设备。检查 SSH IP、端口、账户、密码、eNSP/网络连通性与目标设备的 SSH 服务。"),
        ("任务停止", "停止尚未开始的后续节点；已发送到设备的命令无法从客户端“收回”，需要按回显和现场状态处理。"),
    ], [2400, 6960])
    add_h2(doc, "9.3 Undo 与 PC ping 验收")
    add_bullet(doc, "部分已成功的 vlan_access 计划可使用受限 Undo。Undo 只允许作用于当前拓扑中直接连接 PC 的端口；上联和保护端口会被拒绝。使用前仍要阅读将发送的命令。")
    add_bullet(doc, "如 PC 有明确 SSH 授权，可在成功下发后发起 PC ping 验收。系统只使用 Linux/Windows ping 白名单，不能把该入口当作任意远程命令终端。")
    add_bullet(doc, "验证通过和自动 save 并不等同于业务验收。仍应按现场变更规范检查业务连通性、冗余、告警和回退条件。")

    add_h1(doc, "10. 模板管理")
    add_para(doc, "当一次配置任务的拓扑、需求、思路和命令值得复用时，可从“配置规划”保存为模板。模板是本地不可变快照，便于积累经过人工认可的组织方式，而不是把旧 CLI 原样复制到新设备。")
    add_h2(doc, "10.1 保存与使用模板")
    add_number(doc, "在已生成设备命令的规划任务中，点击“保存为模板”，填写标题和简介。配置思路为空时不能保存。")
    add_number(doc, "进入“模板管理”查看模板详情：拓扑、配置要求、配置思路和逐设备命令都可回看。")
    add_number(doc, "创建新规划任务时选择一个模板。LLM 会把它当作角色划分、实施顺序和命令组织参考；当前任务的拓扑、手册、端口、VLAN 和地址优先。")
    add_h2(doc, "10.2 模板维护")
    add_bullet(doc, "可编辑标题和简介，便于分类检索；模板快照内容不会被不小心覆盖。")
    add_bullet(doc, "可导出单个模板，也可导入单个模板归档。发生重名时按提示选择是否覆盖。")
    add_bullet(doc, "删除模板只影响模板库，不删除原始拓扑、手册或已执行任务。")

    add_h1(doc, "11. 单项导入、导出与删除")
    add_para(doc, "拓扑、手册和模板都支持“单项”导入导出，不会一次把所有本地数据打包。桌面版导出会打开文件保存对话框，让用户选择目录和文件名；完成后页面会提示实际保存到哪里。")
    add_table(doc, ["对象", "导出包含", "导入与删除"], [
        ("拓扑", "节点、链路、接口标签和可用的当前配置快照；不包含 SSH 密码。", "同名时确认覆盖；删除当前保存拓扑前先确认是否仍需要它创建新任务。"),
        ("手册", "原文、抽取结果、命令、辅助映射和已构建向量。", "导入后可直接恢复检索能力；删除后不能用于新任务。"),
        ("模板", "标题、简介、拓扑快照、需求、思路和命令。", "导入时确认重名覆盖；删除不影响原始任务。"),
    ], [1600, 4000, 3760])
    add_note(doc, "归档安全", "归档用于本机数据迁移或团队内受控共享。导出前仍应检查是否包含内部拓扑、设备 IP、命令和手册版权内容；不应直接发送到不受控的公开位置。", CAUTION)

    add_h1(doc, "12. 常见问题")
    add_table(doc, ["现象", "先检查什么", "建议处理"], [
        ("导入手册没有动静", "当前导入任务状态与详细信息。", "刷新手册管理页，确认任务是否 queued/running。CHM 再检查 7-Zip 是否可用；failed 时从断点重试。"),
        ("CHM 解析失败", "文件是否完整、7z.exe 是否在 PATH。", "安装或修复 7-Zip，避免用 hh.exe 解包；重新发起或断点重试。"),
        ("LLM 连不上", "Base URL、Model、API Key、公司网络。", "点击“测试 LLM 连接”；先用 thinking=off 排查端点兼容性，再尝试 adaptive。"),
        ("Embedding 建不起来", "Embedding URL、模型、维度、批量大小。", "将批量调小到 1 或 2；确认端点接受 input 数组及 dimensions（或留空维度）。"),
        ("拓扑保存失败", "拓扑名称、浏览器/桌面窗口状态和日志。", "重试保存；确认每个节点字段独立填写，必要时导出当前项或查看本地日志。"),
        ("前往配置规划无响应", "拓扑是否已保存、手册是否已完成抽取。", "重新选择已保存拓扑，在规划页手工选择拓扑和手册后继续。"),
        ("命令证据不够", "需求是否太笼统、Embedding 是否已建。", "写清目标、端口和约束；查看证据分页，必要时手动查手册并编辑草案。"),
        ("EXE 打开未响应", "完整发布目录、WebView2、数据日志。", "不要单独复制 EXE；检查 _internal 是否存在，结束残留进程后重启，查看 data\\logs。"),
    ], [1900, 2600, 4860])

    add_h1(doc, "13. 本地数据与安全")
    add_h2(doc, "13.1 应保留和应保护的内容")
    add_bullet(doc, "SQLite、手册原文、解包结果、向量、日志和导出文件都可能包含设备地址、拓扑和命令信息，应按内部资料管理。")
    add_bullet(doc, "API Key 不应出现在截图、模板、导出文件、README、Git 记录或问题单中。使用 .env.example 放占位符即可。")
    add_bullet(doc, "SSH 密码不应写进手册、模板标题、需求文本或 Git。下发前只在当前设备的安全输入中填写。")
    add_bullet(doc, "从受控环境导出手册时，注意厂商手册版权和组织的数据合规要求。")
    add_h2(doc, "13.2 为什么手册和向量留在本机")
    add_para(doc, "导入管道会把手册抽取结果和 Embedding 向量存入本地 SQLite。Embedding 文本会调用用户配置的兼容接口，因此应确认该模型服务的安全边界；系统本身不依赖云向量库或 GPU。向量索引不可用时仍可使用精确命令匹配与全文检索。")

    add_h1(doc, "14. GitHub 与二次开发")
    add_para(doc, "完整项目可从 GitHub 获取。读者可以下载、阅读实现，并在符合许可证、手册版权和组织安全规范的前提下继续优化。")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("项目地址：")
    set_run_font(r, size=10.5, bold=True, color=INK)
    add_external_hyperlink(p, "https://github.com/chenjinyu0723/network-automation")
    add_h2(doc, "14.1 下载与本地开发")
    add_code(doc, [
        "git clone https://github.com/chenjinyu0723/network-automation.git",
        "cd network-automation",
        "uv sync --extra dev",
        "pnpm install --frozen-lockfile",
    ])
    add_para(doc, "之后按第 3.2 节启动前后端。需要重新构建 Windows 桌面版时：")
    add_code(doc, [
        "uv sync --extra desktop",
        ".\\scripts\\build_desktop.ps1",
    ])
    add_h2(doc, "14.2 建议的优化方向")
    add_bullet(doc, "为新的厂商手册增加解析适配与命令方言验证规则，而不是把能力限制在华为或单一格式。")
    add_bullet(doc, "为高频能力（例如堆叠、Eth-Trunk、冗余路由）增加专用确定性编译器、状态断言和回滚策略，同时保留通用手册驱动路径。")
    add_bullet(doc, "建立更多真实或模拟场景的命令评测集，度量证据命中、命令可执行性、模型延迟与人工修改率。")
    add_bullet(doc, "改进导入质量诊断，例如扫描 PDF 检测、章节结构修复、厂商版本差异对比和失败页可视化。")

    add_h1(doc, "附录 A. Git 忽略规则")
    add_para(doc, "根目录 .gitignore 已包含以下规则。它的目的不是隐藏源代码，而是避免把本机运行数据、真实密钥、导入手册、设备回显和桌面构建产物上传到 GitHub。若组织有意发布 PDF 文档或手册示例，应先确认版权，再把对应文件从忽略规则中显式放行。")
    add_code(doc, [
        "# Python", "__pycache__/", "*.py[cod]", ".venv/", ".pytest_cache/", ".mypy_cache/", "",
        "# Node and generated desktop artifacts", "node_modules/", "apps/web/dist/", "*.tsbuildinfo", "dist/", "build/", "release/", "",
        "# Local runtime data and secrets", "data/", "*.db", "*.sqlite", "*.sqlite3", ".env", ".env.*", "!.env.example", "*.key", "*.pem", "*.p12", "*.pfx", "",
        "# Imported manuals and their extracted copies", "manuals/", "*.chm", "*.pdf", "",
        "# Local diagnostics and user exports", "logs/", "exports/", "qa/", "*.log", "",
        "# IDE / OS", ".idea/", ".vscode/", "Thumbs.db", "Desktop.ini",
    ])
    add_table(doc, ["不应上传", "原因"], [
        ("data/、数据库和日志", "可能包含拓扑、设备地址、手册内容、向量、任务审计和设备回显。"),
        (".env、私钥、证书", "可能包含模型 API Key、SSH 相关凭据或公司证书。仅保留不含真实值的 .env.example。"),
        ("导入的 CHM/PDF/HTML 手册", "可能受厂商版权约束，也可能是内部受控文档。公开仓库只保留说明和可公开的小样本。"),
        ("release、build、dist、node_modules", "可由源代码和锁文件重新生成，上传会增加仓库体积并引入平台差异。"),
    ], [3000, 6360])

    add_h1(doc, "附录 B. 上线前检查清单")
    add_bullet(doc, "已从完整 release\\NetworkAutomation\\ 目录启动，并确认 WebView2、7-Zip 和本机数据目录可用。")
    add_bullet(doc, "LLM 已通过连接测试；Embedding（如使用）已成功完成目标手册的索引。")
    add_bullet(doc, "任务选择了正确的、抽取完成的手册；对命令方言、版本和适用范围已进行人工复核。")
    add_bullet(doc, "拓扑中每条要参与配置的链路已填写两端真实接口，设备 IP/网关/SSH 信息没有从其他节点错误继承。")
    add_bullet(doc, "保护端口、生产上联和不允许改动的接口已明确标记，且不在当前可写命令中。")
    add_bullet(doc, "配置思路已明确设备角色、步骤、接口和验证方式，且已由人工保存确认。")
    add_bullet(doc, "每台设备命令都有足够的手册证据或已被人工查证；“未验证”提示已处理。")
    add_bullet(doc, "已准备现场验证和回退方案；只在授权窗口内逐台确认发送，并在回显中确认 save 结果。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--toc-pages", type=Path)
    args = parser.parse_args()
    toc_pages = {}
    if args.toc_pages and args.toc_pages.exists():
        toc_pages = json.loads(args.toc_pages.read_text(encoding="utf-8"))
    build_document(toc_pages)
    print(OUTPUT)


if __name__ == "__main__":
    main()
