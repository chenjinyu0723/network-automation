from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs") / "工业交换机自动配置用户手册.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "607080"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FC"
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"


def set_run_font(run, size: float, color: str = "1B2738", bold: bool = False, italic: bool = False) -> None:
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:  # type: ignore[no-untyped-def]
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:  # type: ignore[no-untyped-def]
    """Set fixed 9360-DXA table geometry and matching cell widths."""

    table.autofit = False
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for column, width in zip(grid.gridCol_lst, widths, strict=True):
        column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row) -> None:  # type: ignore[no-untyped-def]
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_note_paragraph(paragraph, fill: str, border: str) -> None:  # type: ignore[no-untyped-def]
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)
    borders.append(left)
    p_pr.append(borders)


def add_page_number(paragraph) -> None:  # type: ignore[no-untyped-def]
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("AI Agent 工业交换机自动配置"), 9, MUTED, bold=True)
    header.add_run("   |   本地单用户版操作手册")
    for run in header.runs[1:]:
        set_run_font(run, 9, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer.add_run("用户手册  |  第 ")
    add_page_number(footer)
    footer.add_run(" 页")
    for run in footer.runs:
        set_run_font(run, 9, MUTED)


def add_cover(document: Document) -> None:
    for _ in range(9):
        document.add_paragraph()
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("本地单用户版  |  操作指南"), 10.5, "4E8CA7", bold=True)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("AI Agent 工业交换机\n自动配置用户手册"), 27, "203748", bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(44)
    set_run_font(subtitle.add_run("从手册注入、手绘拓扑到逐台确认下发"), 14, "456070")

    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(5)
    set_run_font(info.add_run("适用范围：Windows 桌面应用 / 本地 FastAPI / React 工作台"), 10.5, MUTED)
    update = document.add_paragraph()
    update.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(update.add_run("版本：0.1  |  更新日期：2026-08-16"), 10.5, MUTED)
    document.add_page_break()


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    set_run_font(paragraph.add_run(text), {1: 16, 2: 13, 3: 12}[level], {1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)


def add_text(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(paragraph.add_run(bold_prefix), 11, "1B2738", bold=True)
        set_run_font(paragraph.add_run(text[len(bold_prefix):]), 11)
    else:
        set_run_font(paragraph.add_run(text), 11)


def add_bullets(document: Document, entries: list[str], numbered: bool = False) -> None:
    for entry in entries:
        paragraph = document.add_paragraph(style="List Number" if numbered else "List Bullet")
        set_run_font(paragraph.add_run(entry), 11)


def add_note(document: Document, title: str, body: str, *, warning: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    style_note_paragraph(paragraph, "FFF5E8" if warning else PALE_BLUE, "D9834A" if warning else "4E8CA7")
    set_run_font(paragraph.add_run(f"{title}  "), 10.5, "A45F27" if warning else DARK_BLUE, bold=True)
    set_run_font(paragraph.add_run(body), 10.5)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers, strict=True):
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        set_run_font(paragraph.add_run(header), 10.5, "203748", bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row, strict=True):
            paragraph = cell.paragraphs[0]
            set_run_font(paragraph.add_run(text), 10.2)
    set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def build_manual() -> None:
    document = Document()
    configure_document(document)
    add_cover(document)

    add_heading(document, "1. 使用前准备", 1)
    add_text(document, "本应用将设备手册转换为可检索的命令知识库，并把用户绘制的拓扑与配置需求编排为逐设备的候选命令。所有数据、手册和执行审计仅保存在本机。")
    add_table(document, ["项目", "要求"], [
        ["启动方式", "保留 release\\NetworkAutomation 目录完整，双击 NetworkAutomation.exe。"],
        ["运行数据", "%LOCALAPPDATA%\\NetworkAutomation\\data，包括 SQLite、原始手册、解包文件、导出和日志。"],
        ["桌面运行时", "Windows 10/11 通常已带 Microsoft Edge WebView2 Runtime；若窗口无法打开，请先安装该运行时。"],
        ["CHM 导入", "安装 7-Zip，并保证 7z.exe 可被系统找到。应用不会使用 hh.exe 解包 CHM。"],
        ["网络安全", "仅在已授权网络中填写 SSH 凭据。系统不保存 SSH 密码；实际下发前仍需逐台确认。"],
    ], [2200, 7160])
    add_note(document, "安全边界", "计划、检索和命令审查可调用 LLM；保护端口、静态校验、Netmiko 下发、验证和 save 由确定性规则控制。", warning=True)

    add_heading(document, "2. 快速工作流", 1)
    add_bullets(document, [
        "在“设置”页配置 LLM 与 Embedding 的 OpenAI 兼容接口，并使用“测试 LLM 连接”验证连通性。",
        "在“手册管理”页导入 CHM、PDF、HTML 或文本手册，等待命令知识库抽取完成。",
        "在“拓扑编辑”页拖入交换机和 PC，连线、填真实端口、设备地址与交换机 SSH 信息，然后保存 revision。",
        "在“配置规划”页输入需求，审阅每台设备的意图、手册证据和候选命令，逐台批准。",
        "在“下发与结果”页逐台执行、查看校验结果；校验通过后系统才会自动 save。",
    ], numbered=True)

    add_heading(document, "3. 设置 LLM 与 Embedding", 1)
    add_text(document, "打开“设置”页，分别填写 LLM 和 Embedding 的 Base URL、API Key、模型名与温度。接口必须兼容 OpenAI 请求格式。API Key 仅存入本机系统凭据存储，不写入 SQLite 或前端状态。")
    add_table(document, ["字段", "使用建议"], [
        ["LLM Base URL / Model", "填写私有模型或内网网关的 OpenAI 兼容地址和聊天模型名。"],
        ["Temperature", "生产配置建议 0 到 0.3，优先稳定性。"],
        ["推理策略", "选择“自适应（推荐）”：需求理解、检索判断、命令计划和审查会开启 thinking；静态校验与下发不调用 LLM。"],
        ["Embedding", "配置后可构建本地向量索引；未配置时仍使用命令名精确匹配和 SQLite FTS5。"],
    ], [2200, 7160])

    add_heading(document, "4. 导入与管理手册", 1)
    add_text(document, "在“手册管理”页将文件拖入导入区。CHM 会先由本机 7-Zip 解包为 HTML，再解析目录、命令页、参数、视图、约束、示例与适用范围。PDF、HTML、TXT 和 Markdown 使用各自适配器进入同一知识结构。")
    add_heading(document, "4.1 查看进度和处理失败", 2)
    add_bullets(document, [
        "导入页会显示最近任务的状态、当前阶段、页数进度和详细信息；重新打开应用后仍会恢复并继续轮询。",
        "CHM 的典型阶段为 queued、extract_chm、parse_html、completed。大手册的 HTML 页面解析需要较长时间。",
        "失败时点击“从断点重试”。已持久化的页面不会重复入库，任务会从已提交位置继续。",
        "若提示找不到 7-Zip，请安装 7-Zip 后再次重试；不要依赖 Windows hh.exe 的解包结果。",
    ])
    add_heading(document, "4.2 构建检索增强", 2)
    add_text(document, "当手册状态为 completed 或 completed_with_issues 时，可点击“构建 Embedding”。该任务在本机后台运行，向量以 SQLite BLOB 保存并由 CPU 计算余弦相似度，不依赖云向量数据库。")

    add_heading(document, "5. 绘制拓扑", 1)
    add_heading(document, "5.1 添加节点和填写属性", 2)
    add_text(document, "从左侧设备面板添加“交换机”或“PC 终端”。单击节点后，在右侧填写名称、IP、掩码前缀和网关；交换机还需填写 SSH IP、端口、用户名及受保护端口。PC 默认不支持 SSH。")
    add_table(document, ["节点", "可连接端口", "实际含义"], [
        ["交换机", "顶部上 1、上 2；底部下 1、下 2", "上联端口作为连线目标，下联端口作为连线起点。端口编号只表示画布连接位置，不会自动变成 GE 命令。"],
        ["PC", "顶部端口", "单一终端网口，通常从交换机下联端口连接到该端口。"],
    ], [1800, 2500, 5060])
    add_note(document, "端口填写规则", "画布的“上 1/下 1”等只是交互端点。请在“链路端口映射”中填写设备真实接口，例如交换机填 GE0/0/1，PC 填 Ethernet0/0/1。系统会保留你填写的 GE 或 GigabitEthernet 形式，不会擅自改写。", warning=True)
    add_heading(document, "5.2 连线与删除", 2)
    add_bullets(document, [
        "从交换机底部的下 1 或下 2 圆点拖线，连接到 PC 顶部端口或另一台交换机顶部的上 1 或上 2 圆点。",
        "在链路端口映射区域分别填写两端的真实接口名。未映射端口不会被系统猜测，也不会被纳入自动配置。",
        "右键单击交换机或 PC，选择“删除设备及关联链路”；右键单击连线，选择“删除链路”。",
        "删除交换机或 PC 时，所有与该节点相连的链路会同时清除；删除操作仅影响尚未保存的当前画布。",
        "完成后点击“保存拓扑 revision”。配置任务引用的是不可变 revision；修改拓扑后必须再次保存。",
    ], numbered=True)

    add_heading(document, "6. 生成配置计划", 1)
    add_text(document, "进入“配置规划”页，选择已完成的手册与保存的拓扑 revision，输入业务需求。系统通过 LangGraph 按顺序执行输入校验、可选 LLM 意图精炼、混合检索、受证据命令编译、静态校验和人工审批。")
    add_table(document, ["界面内容", "需要检查的内容"], [
        ["需求 / 规划思路", "确认 VLAN、接入口和连通性目标与业务需求一致。"],
        ["所选手册 / 现场信息", "确认所选手册正确；现场型号与版本只用于审计，不影响规划门禁。"],
        ["手册命令证据", "每条写命令都应有命令页证据、语法和适用范围。"],
        ["生成命令", "确认仅包含拓扑定义的端口。云口、上联口和受保护端口不能出现在写命令中。"],
    ], [2200, 7160])
    add_note(document, "命令安全门", "LLM 只能输出结构化意图、检索词或受限命令计划，不能自由编写后直接下发 CLI。没有已完成的所选手册、手册证据、拓扑端口或静态校验时，系统将阻断规划。", warning=True)

    add_heading(document, "7. 逐台下发与验证", 1)
    add_bullets(document, [
        "在“下发与结果”页选择本地配置任务和设备计划。SSH 地址、端口和用户名来自被冻结的拓扑快照，不需要手工复制任务 ID。",
        "逐台确认发送。系统会在写入前记录现场 display version 供审计，并执行手册证据、审批 revision、保护端口、前快照和连接信息检查。",
        "仅当命令执行与验证均通过时才自动保存配置。验证失败时不执行 save，并保留命令输出与错误记录。",
        "如已授权 PC 的 SSH，可对已完成执行记录发起 PC ping 验收；仅使用 Linux/Windows ping 白名单。",
    ], numbered=True)

    add_heading(document, "8. 常见问题", 1)
    add_table(document, ["现象", "处理方式"], [
        ["导入后一直没有进度", "确认运行的是最新 NetworkAutomation.exe；在手册管理查看任务状态。若为 failed，先看详情并点击断点重试。"],
        ["CHM 无法导入", "确认文件扩展名为 .chm，安装 7-Zip；查看任务详情中的解包错误。"],
        ["拓扑无法连线", "必须从交换机下方端口圆点开始，到交换机顶部或 PC 顶部端口圆点结束。"],
        ["规划无法生成命令", "确认选择了已完成的正确手册，并检查命令证据、链路端口映射和静态校验提示。"],
        ["命令没有覆盖预期端口", "检查链路端口映射是否写入真实接口；上联、云口和受保护端口会被明确排除。"],
        ["LLM 测试失败", "检查 Base URL、模型名和 API Key；如端点不支持 thinking，系统会自动关闭 thinking 重试并记录降级。"],
        ["下发后没有 save", "这是验证失败或前置安全检查失败的保护结果；查看执行输出后修正计划并重新批准。"],
    ], [2600, 6760])

    add_heading(document, "9. 操作验收清单", 1)
    add_bullets(document, [
        "手册任务已完成，命令数非零，必要时已建立 Embedding 索引。",
        "规划任务已选择正确且完成抽取的手册；现场型号/版本可选做只读审计。",
        "拓扑中每条链路已填写两端真实接口；交换机到 PC 的接入口明确，受保护端口已填写。",
        "每台设备计划有手册证据、明确端口和静态校验结果。",
        "每台设备均已单独审批，执行后查看验证与 save 状态；必要时完成 PC ping 验收。",
    ])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_manual()
